#=======================================================
#Daniel 2024-03-19
#=======================================================
#IMPORTS
import numpy as np
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches
import os

#=======================================================
#PROCEDIMENTO

class Calculador:

    def __init__(self):
        pass

    def processa(self, img, D):

        matriz = np.asarray(img)
        img_largura, img_altura = img.size

        #primeira linha
        for i in range(img_altura):
            if np.any(matriz[i] != 0):
                linha_ini = i
                break

        #ultima linha
        for i in range(img_altura-1,0,-1):
            if np.any(matriz[i] != 0):
                linha_fim = i
                break

        #dimensões
        n_linhas = linha_fim - linha_ini + 1
        b = h = D / n_linhas

        #centro de gravidade
        soma_Ay = 0
        soma_A = 0
        for i in range(img_altura):
            for j in range(img_largura):
                if matriz[i][j] > 0:
                    A = b * h
                    y_linha = (linha_fim-i+1)*h - h/2
                    soma_Ay += A * y_linha
                    soma_A += A

        Y_linha = soma_Ay/soma_A

        #momento de inercia
        I = 0
        for i in range(img_altura):
            for j in range(img_largura):
                if matriz[i][j] > 0:
                    d = (linha_fim-i+1)*h - h/2 - Y_linha
                    I += b*h**3/12 + b*h*d**2

        #visualização da imagem
        posicao_linha = (h*linha_fim + h/2 - Y_linha) / h

        return img_largura, img_altura, Y_linha, I, posicao_linha
    
class Word:

    def __init__(self):
        pass
    
    def converte(self, img_largura, img_altura, Y_linha, I, D):
        relatorio = Document()

        relatorio.add_heading("Relatório - Momento de Inércia", 0)
        p1 = relatorio.add_paragraph("O presente documento tem como objetivo relatar a entrega dos resultados obtidos por meio de um programa de cálculo de momento de inércia de uma seção. Este programa, desenvolvido com o propósito de analisar a distribuição de massa e a resistência à rotação de objetos em diversas aplicações, oferece uma ferramenta essencial para engenheiros, projetistas e pesquisadores que necessitam compreender e otimizar o comportamento estrutural de componentes e sistemas.")
        p1.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        p1.style.font.name = 'Times New Roman'
        p1.style.font.size = Pt(12)
        p1.paragraph_format.space_before = Pt(12)

        p2 = relatorio.add_paragraph("A seguir, segue a imagem da seção calculada:")
        p2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        p2.style.font.name = 'Times New Roman'
        p2.style.font.size = Pt(12)
        p2.paragraph_format.space_before = Pt(12)

        img_para = relatorio.add_paragraph()
        img_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = img_para.add_run()
        run.add_picture("imagem.png", width=Inches(4.0), height=Inches(4.0))

        p3 = relatorio.add_paragraph("Foto 1 - Seção adotada")
        p3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p3.style.font.name = 'Times New Roman'
        p3.style.font.size = Pt(12)
        p3.paragraph_format.space_before = Pt(2)

        p4 = relatorio.add_paragraph("A seguir, a tabela com os resultados mais importantes, juntamente com a posição da linha do CG:")
        p4.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        p4.style.font.name = 'Times New Roman'
        p4.style.font.size = Pt(12)
        p4.paragraph_format.space_before = Pt(12)

        # Adicionando a tabela
        table = relatorio.add_table(rows=6, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Mesclando as células da linha superior
        top_row = table.rows[0]
        top_row.cells[0].merge(top_row.cells[1])
        top_row.cells[0].text = "Dados da Seção"

        # Preenchendo o restante da tabela
        table.cell(1, 0).text = "Largura da imagem (pixels)"
        table.cell(1, 1).text = str(img_largura)
        table.cell(2, 0).text = "Altura da imagem (pixels)"
        table.cell(2, 1).text = str(img_altura)
        table.cell(3, 0).text = "Altura da imagem (mm)"
        table.cell(3, 1).text = str(D)
        table.cell(4, 0).text = "Altura da linha do CG (mm)"
        table.cell(4, 1).text = str(Y_linha)
        table.cell(5, 0).text = "Momento de inércia (mm^4)"
        table.cell(5, 1).text = str(I)

        p5 = relatorio.add_paragraph("Por fim, segue a imagem da seção calculada, com sua linha de CG:")
        p5.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        p5.style.font.name = 'Times New Roman'
        p5.style.font.size = Pt(12)
        p5.paragraph_format.space_before = Pt(12)

        img_para = relatorio.add_paragraph()
        img_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = img_para.add_run()
        run.add_picture("fig.png", width=Inches(4.0), height=Inches(4.0))

        p6 = relatorio.add_paragraph("Foto 2 - Seção adotada")
        p6.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p6.style.font.name = 'Times New Roman'
        p6.style.font.size = Pt(12)
        p6.paragraph_format.space_before = Pt(2)

        os.remove("imagem.png")
        os.remove("fig.png")

        relatorio.save("relatorio.docx")